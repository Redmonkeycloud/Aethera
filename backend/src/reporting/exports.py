"""Export report to various formats (Docx, PDF, Excel, CSV)."""

from __future__ import annotations

import csv
from pathlib import Path
from typing import Any, Dict

from ..logging_utils import get_logger

logger = get_logger(__name__)


class ReportExporter:
    """Export reports to various formats."""

    @staticmethod
    def to_docx(content: str, output_path: Path, metadata: Dict[str, Any] | None = None) -> Path:
        """
        Export report to Microsoft Word (.docx) format.

        Args:
            content: Markdown or plain text content
            output_path: Output file path
            metadata: Optional document metadata (title, author, etc.)

        Returns:
            Path to generated file
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re

            doc = Document()

            # Set metadata
            if metadata:
                core_props = doc.core_properties
                if "title" in metadata:
                    core_props.title = str(metadata["title"])
                if "author" in metadata:
                    core_props.author = str(metadata["author"])
                if "subject" in metadata:
                    core_props.subject = str(metadata["subject"])

            # Parse markdown-like content
            lines = content.split("\n")
            current_paragraph = None

            for line in lines:
                line = line.strip()

                if not line:
                    if current_paragraph:
                        current_paragraph = None
                    continue

                # Headers
                if line.startswith("# "):
                    heading = doc.add_heading(line[2:], level=1)
                    current_paragraph = None
                elif line.startswith("## "):
                    heading = doc.add_heading(line[3:], level=2)
                    current_paragraph = None
                elif line.startswith("### "):
                    heading = doc.add_heading(line[4:], level=3)
                    current_paragraph = None
                elif line.startswith("#### "):
                    heading = doc.add_heading(line[5:], level=4)
                    current_paragraph = None
                # Lists
                elif line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(line[2:], style="List Bullet")
                    current_paragraph = p
                elif line.startswith("1. ") or re.match(r"^\d+\. ", line):
                    p = doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
                    current_paragraph = p
                # Bold/italic (simple)
                else:
                    p = doc.add_paragraph(line)
                    current_paragraph = p

            doc.save(str(output_path))
            logger.info("Exported DOCX to %s", output_path)
            return output_path

        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

    @staticmethod
    def to_pdf(content: str, output_path: Path, metadata: Dict[str, Any] | None = None) -> Path:
        """
        Export report to PDF format.

        Args:
            content: Markdown or HTML content
            output_path: Output file path
            metadata: Optional document metadata

        Returns:
            Path to generated file
        """
        try:
            from weasyprint import HTML
            try:
                from markdown import markdown
            except ImportError:
                # Fallback: use basic markdown-like parsing
                def markdown(text: str, **kwargs) -> str:
                    # Simple markdown to HTML conversion
                    html = text.replace("\n", "<br>\n")
                    html = html.replace("# ", "<h1>").replace("\n", "</h1>\n", 1) if "# " in html else html
                    return html

            # Convert markdown to HTML
            html_content = markdown(content, extensions=["extra", "tables"])

            # Add basic styling
            html_doc = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    @page {{
                        size: A4;
                        margin: 2cm;
                    }}
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                    }}
                    h1 {{ font-size: 24pt; margin-top: 20pt; }}
                    h2 {{ font-size: 20pt; margin-top: 16pt; }}
                    h3 {{ font-size: 16pt; margin-top: 12pt; }}
                    p {{ margin: 8pt 0; }}
                    ul, ol {{ margin: 8pt 0; padding-left: 20pt; }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 12pt 0;
                    }}
                    th, td {{
                        border: 1px solid #ddd;
                        padding: 8pt;
                        text-align: left;
                    }}
                    th {{
                        background-color: #f2f2f2;
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """

            HTML(string=html_doc).write_pdf(output_path)
            logger.info("Exported PDF to %s", output_path)
            return output_path

        except ImportError:
            raise ImportError("weasyprint not installed. Install with: pip install weasyprint markdown")

    @staticmethod
    def to_excel(
        data: Dict[str, Any],
        output_path: Path,
        sheet_name: str = "Report Data",
    ) -> Path:
        """
        Export structured data to Excel format.

        Args:
            data: Dictionary of data tables (sheet_name -> list of dicts)
            output_path: Output file path
            sheet_name: Default sheet name

        Returns:
            Path to generated file
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment

            wb = Workbook()
            wb.remove(wb.active)  # Remove default sheet

            for sheet_name, rows in data.items():
                ws = wb.create_sheet(title=sheet_name)

                if not rows:
                    continue

                # Headers
                headers = list(rows[0].keys())
                for col_idx, header in enumerate(headers, start=1):
                    cell = ws.cell(row=1, column=col_idx, value=header)
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal="center")

                # Data rows
                for row_idx, row_data in enumerate(rows, start=2):
                    for col_idx, header in enumerate(headers, start=1):
                        value = row_data.get(header, "")
                        ws.cell(row=row_idx, column=col_idx, value=value)

                # Auto-adjust column widths
                for col in ws.columns:
                    max_length = 0
                    col_letter = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except Exception:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[col_letter].width = adjusted_width

            wb.save(str(output_path))
            logger.info("Exported Excel to %s", output_path)
            return output_path

        except ImportError:
            raise ImportError("openpyxl not installed. Install with: pip install openpyxl")

    @staticmethod
    def to_csv(data: list[Dict[str, Any]], output_path: Path) -> Path:
        """
        Export data to CSV format.

        Args:
            data: List of dictionaries (rows)
            output_path: Output file path

        Returns:
            Path to generated file
        """
        if not data:
            raise ValueError("No data to export")

        with open(output_path, "w", newline="", encoding="utf-8") as f:
            fieldnames = list(data[0].keys())
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(data)

        logger.info("Exported CSV to %s", output_path)
        return output_path

